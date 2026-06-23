"""Testes para app.importer."""

import shutil
from pathlib import Path

import pytest
from PIL import Image

from app.importer import slugify, classify_file, resize_image, convert_docx_to_md, import_game


class TestSlugify:
    def test_basic(self):
        assert slugify("Jogo da Anatomia") == "jogo-da-anatomia"

    def test_accents(self):
        assert slugify("Histologia é legal") == "histologia-e-legal"

    def test_special_chars(self):
        assert slugify("Célula & Tecido!!") == "celula-tecido"

    def test_multiple_hyphens(self):
        assert slugify("A---B") == "a-b"


class TestClassifyFile:
    def test_descricao_docx(self):
        assert classify_file("descricao.docx") == "descricao"

    def test_manual(self):
        assert classify_file("Manual parte 1.jpg") == "manual"

    def test_perfil(self):
        assert classify_file("foto perfil.jpg") == "perfil"

    def test_componentes(self):
        assert classify_file("componentes.jpg") == "componentes"

    def test_caixa(self):
        assert classify_file("caixa do jogo.png") == "componentes"

    def test_unknown(self):
        assert classify_file("notacao.txt") is None


class TestResizeImage:
    def test_resize_large(self, tmp_path):
        src = tmp_path / "large.jpg"
        dst = tmp_path / "resized.jpg"
        img = Image.new("RGB", (3000, 2000), color="red")
        img.save(src)
        resize_image(src, dst)
        with Image.open(dst) as out:
            assert max(out.size) <= 1600

    def test_resize_small_stays(self, tmp_path):
        src = tmp_path / "small.jpg"
        dst = tmp_path / "small_out.jpg"
        img = Image.new("RGB", (800, 600), color="blue")
        img.save(src)
        resize_image(src, dst)
        with Image.open(dst) as out:
            assert out.size == (800, 600)


class TestConvertDocx:
    def test_simple_docx(self, tmp_path):
        from docx import Document
        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Primeiro parágrafo")
        doc.add_paragraph("Segundo parágrafo")
        doc.save(str(docx_path))
        md = convert_docx_to_md(docx_path)
        assert "Primeiro parágrafo" in md
        assert "Segundo parágrafo" in md


class TestImportGame:
    def test_import_simple(self, app, tmp_path, tmp_data_dir):
        origem = tmp_path / "origem"
        origem.mkdir()
        img = Image.new("RGB", (100, 100), color="green")
        img.save(origem / "perfil.jpg")
        from docx import Document
        doc = Document()
        doc.add_paragraph("Regras do jogo")
        doc.save(str(origem / "descricao.docx"))
        with app.app_context():
            gid = import_game(origem, "anatomia", "Jogo Teste", tmp_data_dir)
            from app.models import get_game
            game = get_game(gid)
            assert game is not None
            assert game["nome"] == "Jogo Teste"
            assert game["imagem_perfil"] is not None
            assert game["descricao"] is not None
